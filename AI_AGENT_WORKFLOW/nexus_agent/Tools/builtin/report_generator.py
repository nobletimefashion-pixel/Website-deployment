# Tools/builtin/report_generator.py
"""
Universal Report Generator Tool
────────────────────────────────
The agent describes WHAT it wants in the report (topic, sections, raw content,
web-research notes, a resume, a findings summary — anything) and this tool
turns it into a polished PDF or DOCX file with:

  • Auto-picked or user-specified format  (pdf | docx | both)
  • Multiple layout templates             (professional | academic | minimal | dark | resume)
  • Rich structure: title page, TOC stub, section headings, body text, tables, bullet lists
  • Automatic text wrapping & pagination
  • Page numbers + header/footer
  • No external API calls — 100 % local using reportlab + python-docx

Typical agent prompts that route here:
  "Create a PDF resume for John Doe …"
  "Research Harvard University and write a PDF about it"
  "Turn these findings into a professional DOCX report"
  "Make a two-page security summary as PDF"
"""

from __future__ import annotations

import io
import textwrap
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Literal

from pydantic import BaseModel, Field as PField

from nexus_agent.Tools.base import Tool, ToolInvokation, ToolKind, ToolResult


# ─────────────────────────────────────────────────────────────────────────────
# CONTENT MODEL  (what the agent passes in)
# ─────────────────────────────────────────────────────────────────────────────

class ReportSection(BaseModel):
    heading: str = PField(..., description="Section heading / title")
    body: str     = PField("",  description="Plain-text body for this section")
    bullets: list[str] = PField(
        default_factory=list,
        description="Bullet-point items (shown after body paragraph)"
    )
    table: list[list[str]] = PField(
        default_factory=list,
        description=(
            "Optional table. First row = header row. "
            "Each inner list is one row of cells."
        )
    )
    subsections: list["ReportSection"] = PField(
        default_factory=list,
        description="Optional nested sub-sections (one level deep)"
    )


ReportSection.model_rebuild()   # needed for self-reference


class ReportGeneratorParams(BaseModel):
    title: str = PField(..., description="Report / document title")
    subtitle: str = PField("", description="Optional subtitle or tagline")
    author: str   = PField("", description="Author name (shown on title page / footer)")
    date: str     = PField("", description="Date string — defaults to today if blank")

    sections: list[ReportSection] = PField(
        default_factory=list,
        description=(
            "Ordered list of sections. Each section has a heading, body text, "
            "optional bullet list, optional table, and optional sub-sections."
        )
    )

    format: Literal["pdf", "docx", "both"] = PField(
        "pdf",
        description="Output format: 'pdf', 'docx', or 'both'"
    )
    template: Literal["professional", "academic", "minimal", "dark", "resume"] = PField(
        "professional",
        description=(
            "Visual template: "
            "'professional' (navy + grey), "
            "'academic'     (black + serif), "
            "'minimal'      (clean all-white), "
            "'dark'         (dark bg, light text — PDF only), "
            "'resume'       (compact, ATS-friendly)"
        )
    )
    output_path: str = PField(
        "report",
        description=(
            "Output filename WITHOUT extension. "
            "The tool appends .pdf / .docx automatically. "
            "Default: 'report'"
        )
    )
    page_size: Literal["A4", "Letter"] = PField(
        "A4",
        description="Page size: A4 or Letter"
    )
    include_toc: bool = PField(
        True,
        description="Include a Table of Contents page (PDF only)"
    )
    include_title_page: bool = PField(
        True,
        description="Include a dedicated title / cover page"
    )


# ─────────────────────────────────────────────────────────────────────────────
# COLOUR PALETTES PER TEMPLATE
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class Palette:
    bg:           tuple        # page background  (R,G,B) 0-1
    text:         tuple        # body text
    heading1:     tuple        # H1
    heading2:     tuple        # H2
    heading3:     tuple        # H3
    accent:       tuple        # rule lines, table header bg
    accent_text:  tuple        # text on accent bg
    muted:        tuple        # footer / watermark


_PALETTES: dict[str, Palette] = {
    "professional": Palette(
        bg=(1,1,1), text=(0.13,0.13,0.13),
        heading1=(0.08,0.18,0.40), heading2=(0.12,0.30,0.58),
        heading3=(0.25,0.45,0.70), accent=(0.08,0.18,0.40),
        accent_text=(1,1,1), muted=(0.55,0.55,0.55),
    ),
    "academic": Palette(
        bg=(1,1,1), text=(0,0,0),
        heading1=(0,0,0), heading2=(0.1,0.1,0.1),
        heading3=(0.2,0.2,0.2), accent=(0.2,0.2,0.2),
        accent_text=(1,1,1), muted=(0.5,0.5,0.5),
    ),
    "minimal": Palette(
        bg=(1,1,1), text=(0.15,0.15,0.15),
        heading1=(0,0,0), heading2=(0.2,0.2,0.2),
        heading3=(0.35,0.35,0.35), accent=(0.8,0.8,0.8),
        accent_text=(0,0,0), muted=(0.65,0.65,0.65),
    ),
    "dark": Palette(
        bg=(0.10,0.10,0.12), text=(0.90,0.90,0.90),
        heading1=(0.40,0.80,1.0), heading2=(0.55,0.88,1.0),
        heading3=(0.70,0.92,1.0), accent=(0.25,0.55,0.85),
        accent_text=(1,1,1), muted=(0.55,0.55,0.60),
    ),
    "resume": Palette(
        bg=(1,1,1), text=(0.10,0.10,0.10),
        heading1=(0.06,0.24,0.48), heading2=(0.10,0.35,0.60),
        heading3=(0.20,0.45,0.68), accent=(0.06,0.24,0.48),
        accent_text=(1,1,1), muted=(0.50,0.50,0.50),
    ),
}


# ─────────────────────────────────────────────────────────────────────────────
# MAIN TOOL
# ─────────────────────────────────────────────────────────────────────────────

class ReportGeneratorTool(Tool):
    name = "report_generator"
    description = (
        "Generates a polished PDF and/or DOCX document from structured content. "
        "The agent supplies a title, author, date, and a list of sections (each with "
        "a heading, body text, optional bullet list, optional table, and optional "
        "sub-sections). The tool handles all layout, typography, pagination, "
        "headers/footers, and table-of-contents automatically. "
        "Templates: professional, academic, minimal, dark, resume. "
        "Use this whenever the user asks to 'create a PDF/DOCX', 'write a report', "
        "'build a resume', or 'turn research into a document'."
    )
    kind = ToolKind.WRITE
    schema = ReportGeneratorParams

    # ──────────────────────────────────────────────────────────────────────────
    async def execute(self, invocation: ToolInvokation) -> ToolResult:
        params = ReportGeneratorParams(**invocation.params)

        if not params.date:
            params.date = datetime.utcnow().strftime("%B %d, %Y")

        base = Path(invocation.cwd) / params.output_path
        base.parent.mkdir(parents=True, exist_ok=True)

        created: list[str] = []
        errors:  list[str] = []

        if params.format in ("pdf", "both"):
            try:
                pdf_path = Path(str(base) + ".pdf")
                self._build_pdf(params, pdf_path)
                created.append(str(pdf_path))
            except Exception as exc:
                errors.append(f"PDF: {exc}")

        if params.format in ("docx", "both"):
            try:
                docx_path = Path(str(base) + ".docx")
                self._build_docx(params, docx_path)
                created.append(str(docx_path))
            except Exception as exc:
                errors.append(f"DOCX: {exc}")

        if not created:
            return ToolResult.error_result(
                "Report generation failed:\n" + "\n".join(errors)
            )

        sizes = {p: f"{Path(p).stat().st_size / 1024:.1f} KB" for p in created}
        summary = (
            f"✅ Report generated — {len(created)} file(s)\n"
            + "\n".join(f"  📄 {p}  ({sizes[p]})" for p in created)
            + ("\n\n⚠️  Warnings:\n" + "\n".join(errors) if errors else "")
        )

        return ToolResult.success_result(
            output=summary,
            metadata={
                "files":    created,
                "template": params.template,
                "format":   params.format,
                "sections": len(params.sections),
            }
        )

    # ══════════════════════════════════════════════════════════════════════════
    # PDF ENGINE  (reportlab)
    # ══════════════════════════════════════════════════════════════════════════

    def _build_pdf(self, p: ReportGeneratorParams, out: Path) -> None:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm, mm
        from reportlab.platypus import (
            HRFlowable, Image, PageBreak, Paragraph, SimpleDocTemplate,
            Spacer, Table, TableStyle,
        )
        from reportlab.platypus.tableofcontents import TableOfContents

        pal = _PALETTES[p.template]

        def c(rgb):
            return colors.Color(*rgb)

        page_size = A4 if p.page_size == "A4" else letter
        W, H = page_size

        # ── styles ────────────────────────────────────────────────────────────
        styles = getSampleStyleSheet()

        def mk(name, parent="Normal", **kw) -> ParagraphStyle:
            return ParagraphStyle(name, parent=styles[parent], **kw)

        body_font   = "Times-Roman"    if p.template == "academic" else "Helvetica"
        bold_font   = "Times-Bold"     if p.template == "academic" else "Helvetica-Bold"
        italic_font = "Times-Italic"   if p.template == "academic" else "Helvetica-Oblique"

        ST = {
            "title":    mk("rg_title",    fontSize=28, leading=34, textColor=c(pal.heading1),
                           fontName=bold_font, spaceAfter=6, alignment=1),
            "subtitle": mk("rg_subtitle", fontSize=15, leading=20, textColor=c(pal.heading2),
                           fontName=italic_font, spaceAfter=4, alignment=1),
            "author":   mk("rg_author",   fontSize=11, leading=14, textColor=c(pal.muted),
                           fontName=body_font, spaceAfter=2, alignment=1),
            "date":     mk("rg_date",     fontSize=10, leading=13, textColor=c(pal.muted),
                           fontName=italic_font, alignment=1),
            "h1":       mk("rg_h1",       fontSize=16, leading=20, textColor=c(pal.heading1),
                           fontName=bold_font, spaceBefore=14, spaceAfter=4),
            "h2":       mk("rg_h2",       fontSize=13, leading=17, textColor=c(pal.heading2),
                           fontName=bold_font, spaceBefore=10, spaceAfter=3),
            "h3":       mk("rg_h3",       fontSize=11, leading=15, textColor=c(pal.heading3),
                           fontName=bold_font, spaceBefore=7,  spaceAfter=2),
            "body":     mk("rg_body",     fontSize=10, leading=15, textColor=c(pal.text),
                           fontName=body_font,  spaceAfter=6),
            "bullet":   mk("rg_bullet",   fontSize=10, leading=14, textColor=c(pal.text),
                           fontName=body_font,  leftIndent=18, firstLineIndent=-12, spaceAfter=3),
            "toc_h":    mk("rg_toc_h",   fontSize=14, leading=18, textColor=c(pal.heading1),
                           fontName=bold_font, spaceBefore=0, spaceAfter=8, alignment=1),
            "toc_item": mk("rg_toc_item", fontSize=10, leading=14, textColor=c(pal.text),
                           fontName=body_font, leftIndent=12, spaceAfter=3),
        }

        # ── header / footer callbacks ─────────────────────────────────────────
        def _on_page(canvas, doc):
            canvas.saveState()
            # background fill (dark template)
            if p.template == "dark":
                canvas.setFillColor(c(pal.bg))
                canvas.rect(0, 0, W, H, fill=1, stroke=0)
            # thin rule at top
            canvas.setStrokeColor(c(pal.accent))
            canvas.setLineWidth(1.5)
            canvas.line(2*cm, H - 1.5*cm, W - 2*cm, H - 1.5*cm)
            # header text
            canvas.setFont(body_font, 8)
            canvas.setFillColor(c(pal.muted))
            canvas.drawString(2*cm, H - 1.3*cm, p.title)
            if p.author:
                canvas.drawRightString(W - 2*cm, H - 1.3*cm, p.author)
            # footer rule
            canvas.setStrokeColor(c(pal.accent))
            canvas.line(2*cm, 1.5*cm, W - 2*cm, 1.5*cm)
            # page number
            canvas.setFont(body_font, 8)
            canvas.setFillColor(c(pal.muted))
            canvas.drawCentredString(W / 2, 1.0*cm, f"— {doc.page} —")
            canvas.restoreState()

        def _on_first_page(canvas, doc):
            canvas.saveState()
            if p.template == "dark":
                canvas.setFillColor(c(pal.bg))
                canvas.rect(0, 0, W, H, fill=1, stroke=0)
            canvas.restoreState()

        # ── document ──────────────────────────────────────────────────────────
        doc = SimpleDocTemplate(
            str(out),
            pagesize=page_size,
            leftMargin=2.5*cm, rightMargin=2.5*cm,
            topMargin=2.5*cm,  bottomMargin=2.5*cm,
            title=p.title,
            author=p.author or "Nexus Agent",
        )

        story = []

        # ── title page ────────────────────────────────────────────────────────
        if p.include_title_page:
            story.append(Spacer(1, H * 0.18))
            # accent bar
            story.append(HRFlowable(
                width="100%", thickness=4,
                color=c(pal.accent), spaceAfter=20
            ))
            story.append(Paragraph(self._esc(p.title), ST["title"]))
            if p.subtitle:
                story.append(Spacer(1, 6))
                story.append(Paragraph(self._esc(p.subtitle), ST["subtitle"]))
            story.append(HRFlowable(
                width="60%", thickness=1,
                color=c(pal.accent), spaceBefore=14, spaceAfter=14
            ))
            if p.author:
                story.append(Paragraph(self._esc(p.author), ST["author"]))
            story.append(Paragraph(p.date, ST["date"]))
            story.append(PageBreak())

        # ── TOC page ──────────────────────────────────────────────────────────
        if p.include_toc and p.sections:
            story.append(Paragraph("Table of Contents", ST["toc_h"]))
            story.append(HRFlowable(width="100%", thickness=0.5, color=c(pal.accent), spaceAfter=8))
            for i, sec in enumerate(p.sections, 1):
                story.append(Paragraph(f"{i}.  {self._esc(sec.heading)}", ST["toc_item"]))
                for j, sub in enumerate(sec.subsections, 1):
                    story.append(Paragraph(
                        f"&nbsp;&nbsp;&nbsp;&nbsp;{i}.{j}  {self._esc(sub.heading)}",
                        ST["toc_item"]
                    ))
            story.append(PageBreak())

        # ── sections ──────────────────────────────────────────────────────────
        for sec in p.sections:
            story += self._pdf_section(sec, ST, pal, c, HRFlowable, Paragraph, Spacer, Table, TableStyle, colors, depth=1)

        # ── build ──────────────────────────────────────────────────────────────
        doc.build(
            story,
            onFirstPage=_on_first_page,
            onLaterPages=_on_page,
        )

    def _pdf_section(self, sec, ST, pal, c, HRFlowable, Paragraph, Spacer, Table, TableStyle, colors, depth=1):
        from reportlab.lib.units import cm
        items = []

        h_key = {1: "h1", 2: "h2", 3: "h3"}.get(depth, "h3")
        items.append(Paragraph(self._esc(sec.heading), ST[h_key]))

        # horizontal rule under H1
        if depth == 1:
            items.append(HRFlowable(
                width="100%", thickness=0.8,
                color=c(pal.accent), spaceAfter=4
            ))

        if sec.body:
            for para in sec.body.split("\n\n"):
                para = para.strip()
                if para:
                    items.append(Paragraph(self._esc(para), ST["body"]))

        for bullet in sec.bullets:
            items.append(Paragraph(f"• &nbsp;{self._esc(bullet)}", ST["bullet"]))

        if sec.table:
            items += self._pdf_table(sec.table, ST, pal, c, Table, TableStyle, colors)

        for sub in sec.subsections:
            items += self._pdf_section(sub, ST, pal, c, HRFlowable, Paragraph, Spacer, Table, TableStyle, colors, depth=depth+1)

        items.append(Spacer(1, 8))
        return items

    def _pdf_table(self, rows, ST, pal, c, Table, TableStyle, colors):
        from reportlab.lib.units import cm

        # Build paragraph cells
        def cell(txt, bold=False):
            from reportlab.platypus import Paragraph as P
            font = "Helvetica-Bold" if bold else "Helvetica"
            st = ST["body"].__class__(
                "cell_" + ("h" if bold else "b"),
                parent=ST["body"],
                fontName=font,
                fontSize=9,
                leading=12,
                textColor=c(pal.accent_text if bold else pal.text),
            )
            return P(self._esc(str(txt)), st)

        data = []
        for i, row in enumerate(rows):
            data.append([cell(v, bold=(i == 0)) for v in row])

        col_count = max(len(r) for r in rows)
        col_w = (16 * cm) / col_count

        t = Table(data, colWidths=[col_w] * col_count, repeatRows=1)
        t.setStyle(TableStyle([
            # header row
            ("BACKGROUND",    (0, 0), (-1, 0),   c(pal.accent)),
            ("TEXTCOLOR",     (0, 0), (-1, 0),   c(pal.accent_text)),
            ("FONTNAME",      (0, 0), (-1, 0),   "Helvetica-Bold"),
            ("FONTSIZE",      (0, 0), (-1, 0),   9),
            # body rows
            ("BACKGROUND",    (0, 1), (-1, -1),  colors.white),
            ("ROWBACKGROUNDS",(0, 1), (-1, -1),
             [colors.Color(0.95, 0.97, 1.0), colors.white]),
            ("FONTNAME",      (0, 1), (-1, -1),  "Helvetica"),
            ("FONTSIZE",      (0, 1), (-1, -1),  9),
            ("TEXTCOLOR",     (0, 1), (-1, -1),  c(pal.text)),
            # grid
            ("GRID",          (0, 0), (-1, -1),  0.4, c(pal.muted)),
            ("VALIGN",        (0, 0), (-1, -1),  "TOP"),
            ("TOPPADDING",    (0, 0), (-1, -1),  4),
            ("BOTTOMPADDING", (0, 0), (-1, -1),  4),
            ("LEFTPADDING",   (0, 0), (-1, -1),  6),
            ("RIGHTPADDING",  (0, 0), (-1, -1),  6),
        ]))

        from reportlab.platypus import Spacer as Sp
        return [Sp(1, 6), t, Sp(1, 6)]

    # ══════════════════════════════════════════════════════════════════════════
    # DOCX ENGINE  (python-docx)
    # ══════════════════════════════════════════════════════════════════════════

    def _build_docx(self, p: ReportGeneratorParams, out: Path) -> None:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor

        pal = _PALETTES[p.template]

        def rgb(tup):
            r, g, b = tup
            return RGBColor(int(r*255), int(g*255), int(b*255))

        doc = Document()

        # ── page margins ──────────────────────────────────────────────────────
        for section in doc.sections:
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin   = Inches(1.15)
            section.right_margin  = Inches(1.15)

        # ── title page ────────────────────────────────────────────────────────
        if p.include_title_page:
            tp = doc.add_paragraph()
            tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = tp.add_run(p.title)
            run.bold      = True
            run.font.size = Pt(28)
            run.font.color.rgb = rgb(pal.heading1)

            if p.subtitle:
                sp = doc.add_paragraph(p.subtitle)
                sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in sp.runs:
                    r.italic = True
                    r.font.size = Pt(14)
                    r.font.color.rgb = rgb(pal.heading2)

            doc.add_paragraph()  # spacer

            if p.author:
                ap = doc.add_paragraph(p.author)
                ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in ap.runs:
                    r.font.size = Pt(11)
                    r.font.color.rgb = rgb(pal.muted)

            dp = doc.add_paragraph(p.date)
            dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in dp.runs:
                r.italic = True
                r.font.size = Pt(10)
                r.font.color.rgb = rgb(pal.muted)

            doc.add_page_break()

        # ── TOC placeholder ───────────────────────────────────────────────────
        if p.include_toc and p.sections:
            toc_h = doc.add_paragraph("Table of Contents")
            toc_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in toc_h.runs:
                r.bold = True
                r.font.size = Pt(14)
                r.font.color.rgb = rgb(pal.heading1)

            for i, sec in enumerate(p.sections, 1):
                li = doc.add_paragraph(f"{i}.  {sec.heading}", style="List Number")
                li.paragraph_format.left_indent = Inches(0.3)
                for r in li.runs:
                    r.font.size = Pt(10)

            doc.add_page_break()

        # ── sections ──────────────────────────────────────────────────────────
        for sec in p.sections:
            self._docx_section(doc, sec, rgb, pal, Pt, RGBColor, WD_ALIGN_PARAGRAPH, Inches, depth=1)

        # ── footer with page numbers ──────────────────────────────────────────
        for section in doc.sections:
            footer = section.footer
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = fp.add_run()
            run.font.size = Pt(8)
            run.font.color.rgb = rgb(pal.muted)
            # Add page number field
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), "begin")
            run._r.append(fld)
            instr = OxmlElement("w:instrText")
            instr.text = " PAGE "
            run._r.append(instr)
            fld2 = OxmlElement("w:fldChar")
            fld2.set(qn("w:fldCharType"), "end")
            run._r.append(fld2)

            if p.author:
                fp.add_run(f"  |  {p.author}").font.size = Pt(8)

        doc.save(str(out))

    def _docx_section(self, doc, sec, rgb, pal, Pt, RGBColor, WD_ALIGN_PARAGRAPH, Inches, depth=1):
        style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
        color_map = {1: pal.heading1, 2: pal.heading2, 3: pal.heading3}

        hpara = doc.add_paragraph(sec.heading, style=style_map.get(depth, "Heading 3"))
        for run in hpara.runs:
            run.font.color.rgb = rgb(color_map.get(depth, pal.heading3))
            run.font.size = Pt({1: 15, 2: 13, 3: 11}.get(depth, 11))

        if sec.body:
            for para in sec.body.split("\n\n"):
                para = para.strip()
                if para:
                    bp = doc.add_paragraph(para)
                    for r in bp.runs:
                        r.font.size = Pt(10)
                        r.font.color.rgb = rgb(pal.text)

        for bullet in sec.bullets:
            bl = doc.add_paragraph(bullet, style="List Bullet")
            for r in bl.runs:
                r.font.size = Pt(10)

        if sec.table:
            rows_data = sec.table
            row_count = len(rows_data)
            col_count = max(len(r) for r in rows_data)
            tbl = doc.add_table(rows=row_count, cols=col_count)
            tbl.style = "Table Grid"
            for ri, row in enumerate(rows_data):
                for ci, val in enumerate(row):
                    cell = tbl.cell(ri, ci)
                    cell.text = str(val)
                    run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(str(val))
                    run.font.size = Pt(9)
                    if ri == 0:
                        run.bold = True
                        run.font.color.rgb = rgb(pal.accent_text)
                        # shade header cell
                        from docx.oxml import OxmlElement
                        from docx.oxml.ns import qn
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd = OxmlElement("w:shd")
                        r_, g_, b_ = pal.accent
                        hex_color = f"{int(r_*255):02X}{int(g_*255):02X}{int(b_*255):02X}"
                        shd.set(qn("w:fill"), hex_color)
                        tcPr.append(shd)
                    else:
                        run.font.color.rgb = rgb(pal.text)

        for sub in sec.subsections:
            self._docx_section(doc, sub, rgb, pal, Pt, RGBColor, WD_ALIGN_PARAGRAPH, Inches, depth=depth+1)

    # ──────────────────────────────────────────────────────────────────────────
    # HELPERS
    # ──────────────────────────────────────────────────────────────────────────

    @staticmethod
    def _esc(text: str) -> str:
        """Escape XML special chars for ReportLab Paragraph."""
        return (
            text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace('"', "&quot;")
        )